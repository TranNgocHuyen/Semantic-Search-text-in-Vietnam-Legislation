import os
import csv
import docx
import pandas as pd
# lấy tên, title và content
def get_name_title_content(folder_path, file_name):
    file_path = os.path.join(folder_path, file_name)
    try: 
        document = docx.Document(file_path) # là 1 đối tượng

        if len(document.paragraphs) > 2:
            title = document.paragraphs[1].text+ " " + document.paragraphs[2].text
            full_text = ""
            for i, paragraph in enumerate(document.paragraphs):
                if i>2 and i!= (len(document.paragraphs) - 1):
                    full_text += paragraph.text + "\\n"

            full_text = full_text.strip()#.replace('\n', '\\n')
        else:
            print(file_path)
        return [file_name, title, full_text]
    
    except FileNotFoundError:
        print(f"Error: File '{file_name}' not found.")
        return None





if __name__ == "__main__":
    # tạo bảng
    # with open("dataset_0_500_crawl_doc.csv", 'w', newline='', encoding="utf-8") as file:
    #     writer=csv.writer(file)
    #     field=['FileName', 'TieuDe', 'NoiDungVanBan']
    #     writer.writerow(field)


        folder_path = "D:\\NLP\\process_van_ban\\data_1_00\\"
        folder = os.listdir(folder_path)
        list_data = []
        i=0 # 12,864
        for file in folder:
            i = i +1
            print("======", i)
            data = get_name_title_content(folder_path, file)
            if data is None:
                continue
            else:
                list_data.append(data)
                print(list_data)
            if i==3:
                break
            #print(list_data)
            
        df = pd.DataFrame(list_data, columns=['FileName', 'TieuDe', 'NoiDungVanBan'])
        # converting to CSV file
        df.to_csv("dataset_0_500_crawl_doc_test.csv", encoding =  "utf-8", index = False)

